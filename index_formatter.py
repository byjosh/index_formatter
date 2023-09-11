#  using this needs to function in a venv - the usual /usr/bin/env python3 can misconfigure it
# This file is from
# https://github.com/byjosh/index_formattter
# licensed under https://www.gnu.org/licenses/gpl-3.0.html
import csv
from collections import namedtuple
import html_page_components
import time
import sys
from os import path
import logging
import html
import argparse
# test if we have python-docx available
from importlib import resources
docx_found = 'xxxx'
try:
    docx_found = resources.files('docx')
except ModuleNotFoundError:
    print("ModuleNotFoundError for docx - if you want docx output: create a virtual environment and run: pip install python-docx")

def test_for_docx(docx_found):
    if str(docx_found)[-4:] == 'docx':
        return True
    else:
        return False

logging.info(f'python-docx found for docx export = {test_for_docx(docx_found)}')


# if you want to see the master dictionary created change logging.FATAL to logging.DEBUG
loglevel = logging.FATAL

"""have Location tuple with book and page, then list of tuples consisting of (book, page num, entry)"""
Location = namedtuple('Location', 'book page')
Entry = namedtuple('Entry', 'location notes taglist chapter')
# possible field orders for columns in CSV: Entry, Book, Chapter, Notes, Tags
field_orders = ['ebpnt',
                'epbnt',
                'ebcpnt',
                'enbcpt',
                'enbpt',
                'bpent',
                'bcpent',
                'entbp',
                'entbcp',
                ]
# set some defaults
tagseparator = ": "
outputseparator = ' -- '

# Change docx help message according to availability of python-docx
docx_help = 'Not available on this install currently - but see the README for using a virtual environment and pip install python-docx to allow docx file export (done successfully this help message will change) - .docx files can be used in wordprocessing software'
if test_for_docx(docx_found):
    docx_help = 'you can set this option to export a docx file - docx files can be used in wordprocessing software (congratulations on using pip to enable this option)'

parser = argparse.ArgumentParser(prog="index_as_html.py",
                                 description="takes a CSV file with columns for entry, book #, [chapter#] , page#, description and tags fields (not necessarily in that order and outputs order HTML or a CSV with tags prepended to the entry - if your output has weird characters try saving your input as utf-8 (the format for output) - for docx output: see README and install python-docx using pip")

parser.add_argument('inputCSVfile',
                    help='input CSV file to be processed - see README and help below re: order of columns needed')
parser.add_argument('-o', '--outputtitle', action='store', help="title used in naming output document and in HTML head")
parser.add_argument('-f', '--fields', action='store', choices=field_orders,
                    help=f'specify order of columns (fields) in your CSV file where each letter is initial letter of following (might call your columns slightly differently but this make clear function of each column):\nEntry (what you look up in index e.g. apple),\nBook (or volume e.g. vol 1) , \nChapter [optional], Notes (notes or description on Entry item e.g. popular fruit), \nTags [optional in sense program will check if this is empty column] - comma separated list of classifications that might apply to the entry e.g. for apple depending on context: fruit, tree, cooking, literacy, physics, Abrahamic religion (those are just examples - in botanical context just fruit and tree would be relevant)')
parser.add_argument('-t', '--tags', action='store_true',
                    help=f'process tags so that entries are added with the tag prepended to the entry - using tag separator (defaults to {tagseparator} )')
parser.add_argument('-ts', '--tagseparator', action='store',
                    help=f'The separator between a tag and an entry when adding an entry of tag tagseparator entry default set to {tagseparator} - likely a colon followed by a space')
parser.add_argument('-c', '--csv', action='store_true', help='output file as CSV - if used in conjunctions')
parser.add_argument('-os', '--outputseparator', action='store',
                    help='when outputting HTML the book, chapter, page numbers will be concatenated and joined this seperator is between them')
parser.add_argument('-v', action='store_true', help='logging.INFO level of log verbosity')
parser.add_argument('-vv', action='store_true', help='logging.DEBUG level of logging verbosity')
parser.add_argument('-d','--docx', action='store_true', help=f'{docx_help} ')


args = parser.parse_args()

# set various fields away from defaults
if args.v:
    loglevel = logging.INFO
if args.vv:
    loglevel = logging.DEBUG
logging.basicConfig(level=loglevel)
logging.info(f'CSV dialects are {csv.list_dialects()}')
logging.info(f'arguments passed are {args}')

title = None
order = 'ebpnt'

csv_input_filename = html.escape(args.inputCSVfile)
if args.outputtitle:
    title = html.escape(args.outputtitle)
if args.fields:
    order = html.escape(args.fields)
if args.tagseparator:
    tagseparator = html.escape(args.tagseparator)
if args.outputseparator:
    outputseparator = html.escape(args.outputseparator)
logging.debug(f'sys.argv was {sys.argv}')
logging.info(f'Order is {order} ')


# END reset defaults
# START utility functions
def output_file_message(outputname):
    """Utility function for message indicating successful file writing"""
    print(f'Created {outputname} as a file in the same folder as this script')


def tuple_string(entry_tuple):
    """Function used to get key for ordering book - relies on last char of book sting being a digit"""
    book_num = entry_tuple.location.book.strip()[-1:]
    page_num = entry_tuple.location.page
    if book_num.isdigit() and page_num.isdigit():
        # so book 4 page 40 is not more than book 6 page 7 use x 1000
        float_value = (int(book_num) * 1000) + int(page_num)
        # use a string anyway so results are comparable
        return str(float_value)
    elif not book_num.isdigit() or not page_num.isdigit():
        # it went wrong - use a string
        print(
            f'For better ordering have last character of {entry_tuple.location.book} as number not {book_num} and {page_num} as number')
        return html.escape(f'{book_num}-{page_num}', quote=True)


def not_file(file):
    """Test file is missing"""
    if path.isfile(file):
        return False
    else:
        return True


def not_csv(file):
    """Tests that file misses a .csv extension"""
    if file.split(".")[-1].casefold() == 'csv'.casefold():
        return False
    else:
        return True


# END utility functions

if not_file(csv_input_filename) or not_csv(csv_input_filename):
    # if the input file does not exist exit
    print("first argument does not exist as file or is not a .csv file")
    sys.exit()


def create_master_dict(csv_input_filename: str) -> dict:  #
    """Given a CSV file with columns for entry, book number, page number, notes - in that order - return a dictionary with
    entries as keys and value is list of named tuples of Entry(Location(book,page),notes) format"""
    master_dict = {}
    global order
    with open(csv_input_filename, encoding='utf-8') as csvfile:

        index_reader = csv.reader(csvfile, delimiter=',', quotechar='"')
        entry = order.find('e')
        book = order.find('b')
        chapter = order.find('c')
        page = order.find('p')
        notes = order.find('n')
        tags = order.find('t')
        chapter_no = None
        taglist = None

        for row in index_reader:
            logging.info(f'row length is {len(row)}')
            if chapter != -1:
                chapter_no = row[chapter]
            else:
                chapter_no = None
            if len(row) == len(order):
                taglist = row[tags]
            else:
                taglist = None
            logging.debug(f'Current row is {row}')
            if row[entry] in master_dict.keys():
                master_dict[row[entry]].append(Entry(Location(row[book], row[page]), row[notes], taglist, chapter_no))
            elif row[entry] not in master_dict.keys() and row[entry].strip() != '':
                # the row[entry].strip() != '' means ignore line if no entry
                master_dict[row[entry]] = [Entry(Location(row[book], row[page]), row[notes], taglist, chapter_no)]
            if args.tags is True and taglist != '' and taglist is not None:
                for current_tag in taglist.split(','):
                    current_tag = current_tag.strip()
                    current_tag_entry = f'{current_tag}{tagseparator}{row[entry]}'
                    if current_tag_entry in master_dict.keys():
                        # TODO: this assumes another entry with same tag may exist - valid assumption?
                        master_dict[current_tag_entry].append(
                            Entry(Location(row[book], row[page]), row[notes], taglist, chapter_no))
                    elif current_tag_entry not in master_dict.keys() and row[entry].strip() != '':
                        # the row[entry].strip() != '' means ignore line if no entry
                        master_dict[current_tag_entry] = [
                            Entry(Location(row[book], row[page]), row[notes], taglist, chapter_no)]
        logging.debug(f'dictionary of entries generated is {master_dict}')
        list_keys = list(master_dict.keys())
        last = list_keys[-1]
        midentry = list_keys[int(len(list_keys) / 2)]
        logging.info(f'midpoint sample is {master_dict[midentry]}')
        logging.info(f'endpoint sample is {master_dict[last]}')
    return master_dict


def create_page_html(sorted_keys,entries):
    """Main function creating HTML for output"""

    def html_item_output(item):
        """Creates a div for each Entry tuple per index entry (dict key in master_dict"""
        if not item.chapter:
            chapter = ''
        else:
            chapter = item.chapter + outputseparator
        return f'<div class="moredetails"><span class="location">{html.escape(item.location.book, quote=True)}{html.escape(outputseparator)}{html.escape(chapter)}{html.escape(item.location.page, quote=True)}</span>\
        <span class="notes">{html.escape(item.notes, quote=True)}</span></div>'


    output = html_page_components.get_html_header(title)

    for key in sorted_keys:
        list_of_Entry_tuples = sorted(entries[key], key=tuple_string)
        output += f'<div class="entrycontainer"><div class="entry">{key}</div>'
        for item in list_of_Entry_tuples:
            output += html_item_output(item)
        output += '</div>'

    output += html_page_components.get_html_footer()
    return output


def write_html_file(now,sorted_keys,entries):
    output_name = f'{title}_as_html_{now}.html'
    with open(output_name, mode='w', encoding='utf-8') as file:
        file.write(create_page_html(sorted_keys,entries))
    output_file_message(output_name)
    logging.debug(f'The arguments provided to script were {sys.argv[1:]} as sys.argv and as argparse {args}')


def write_csv_file(now,sorted_keys,entries):
    output_name = f'{title}_as_csv_{now}.csv'
    import csv

    with open(output_name, 'w', encoding='utf-8') as file:
        writer = csv.writer(file, dialect='excel')

        for key in sorted_keys:
            list_of_Entry_tuples = sorted(entries[key], key=tuple_string)
            for item in list_of_Entry_tuples:
                tags = ''
                if item.taglist:
                    tags = item.taglist
                # TODO: tags output to CSV file - so tag: entry is possible but so is tag: tag: tag: entry if one ran the script recursively - desirable or not?
                if item.chapter:
                    writer.writerow([key, item.location.book, item.chapter, item.location.page, item.notes,
                                     tags])
                if not item.chapter:
                    writer.writerow([key, item.location.book, item.location.page, item.notes, tags])

    output_file_message(output_name)

def write_docx_file(now,sorted_keys,entries):
    """Write a docx document - use the alphabet as guide when to insert a header"""
    from docx import Document
    from docx.shared import Mm
    from docx.shared import Pt
    output_name = f'{title}_as_docx_{now}.docx'
    document = Document()
    alphabet = list('ABCDEFGHIJKLMNOPQRSTUVWXYZ')
    prev_char = None

    for key in sorted_keys:
        first_char = f'{key[0]}'.upper()
        if first_char in alphabet and first_char != prev_char:

            document.add_paragraph(first_char,style='Heading 2')
            style = document.styles['Heading 2']
            style.font.size = Pt(48)

            prev_char = first_char
        list_of_Entry_tuples = sorted(entries[key], key=tuple_string)
        p = document.add_paragraph(style='Normal')
        p_format = p.paragraph_format
        p_format.first_line_indent = Mm(-8)
        p.add_run(f'{key} ').bold = True
        for item in list_of_Entry_tuples:
            tags = ''
            if item.taglist:
                tags = item.taglist
            # TODO: tags output to CSV file - so tag: entry is possible but so is tag: tag: tag: entry if one ran the script recursively - desirable or not?
            if item.chapter:
                p.add_run(f'{item.location.book}{outputseparator}{item.chapter}{outputseparator}{item.location.page} ').bold = True
            if not item.chapter:
                p.add_run(
                    f'{item.location.book}{outputseparator}{item.location.page} ').bold = True
            p.add_run(f'{item.notes}\n')
    document.save(output_name)
    output_file_message(output_name)

def write_to_file():
    """writes output to a file"""
    # write the output to a file -
    now = time.strftime('%Y%m%d-%H%M%S')
    entries = create_master_dict(csv_input_filename)
    sorted_keys = sorted(entries.keys(), key=str.casefold)
    if not args.csv and not args.docx:
        logging.info("writing HTML file")
        write_html_file(now,sorted_keys,entries)

    if args.csv:
        logging.info("writing CSV file")
        write_csv_file(now,sorted_keys,entries)
    try:
        if args.docx and test_for_docx(docx_found):
            write_docx_file(now,sorted_keys,entries)
        elif args.docx and test_for_docx(docx_found) is False:
           print("Sorry python-docx module for docx output not found - trying installing it with pip as per instructions in README")
    except Exception as e:
        print("About to raise a docx exception")
        raise e


if __name__ == "__main__":
    write_to_file()
